import os
import re
from datetime import datetime
from docx import Document
from yadisk import YaDisk

import logging
from utils.logging_config import get_logger

from config import YANDEX_DISK_TOKEN

logger = get_logger(__name__)


def extract_grade(file_path):
    """
    Извлекает оценку из последней строки Word документа
    Ожидает формат: '8/10' или '8 / 10'
    """
    try:
        doc = Document(file_path)
        last_line = doc.paragraphs[-1].text.strip()

        # Ищем оценку в формате "число/10"
        match = re.search(r'\d+\s*/\s*10', last_line)
        if match:
            return str(match.group(0))
        return None
    
    except Exception as e:
        error_message = f"Ошибка в структуре документа, нет оценки {e}"
        print(error_message)
        logger.warning(error_message)
        return None


def extract_name_from_text(text):
    """
    Создает безопасное имя файла из первой строки текста
    """
    first_line = text.strip().split('\n')[0].strip()
    # Заменяем все не-буквенные символы на подчеркивания
    name_safe = re.sub(r'[^\w\-]', '_', first_line)
    
    return name_safe


def save_docx(text):
    """
    Сохраняет текст в Word документ на локальный диск
    """
    try:
        # Создаем папку если ее нет
        os.makedirs("output_docs", exist_ok=True)

        # Генерируем имя файла
        name_part = extract_name_from_text(text)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{name_part}_{timestamp}.docx"
        local_path = os.path.join("output_docs", filename)
        
        # Создаем и сохраняем документ
        doc = Document()
        doc.add_paragraph(text)
        doc.save(local_path)

        return(local_path)
    
    except Exception as e:
        error_message = f"Ошибка при сохранении ворд файла {e}"
        print(error_message)
        logger.warning(error_message)
        return None 

def upload_to_yandex(file_path):
    """
    Загружает файл на Яндекс Диск и возвращает публичную ссылку
    """
    try:
        y = YaDisk(token=YANDEX_DISK_TOKEN)
        remote_dir = "/resume_script"

        # Создаем папку на Яндекс Диске если ее нет
        if not y.exists(remote_dir):
            y.mkdir(remote_dir)

        # Загружаем файл
        remote_path = f"{remote_dir}/{os.path.basename(file_path)}"
        y.upload(file_path, remote_path, overwrite=True)

        # Получаем публичную ссылку
        link = y.get_download_link(remote_path)
        return link
    
    except Exception as e:
        error_message = f"Ошибка при загрузке docx на Яндекс диск {e}"
        print(error_message)
        logger.warning(error_message)
        return None


def save_and_upload(text):
    """
    Основная функция: сохраняет текст в DOCX и загружает на Яндекс Диск
    Возвращает (ссылка, оценка)
    """
    try:
        file_path = save_docx(text)
        link = upload_to_yandex(file_path)
        grade = extract_grade(file_path)
        return link, grade
    except Exception as e:
        error_message = f"Ошибка c docx файлом {e}"
        print(error_message)
        logger.warning(error_message)
        return None, None
